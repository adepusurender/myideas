import unittest
import os
import tempfile
import shutil
from unittest.mock import patch, MagicMock, mock_open
from docx import Document
import sys

# Import the module to test
from video_to_subtitles import (
    format_time,
    clean_srt_content,
    save_as_docx,
    batch_translate,
    generate_srt,
    process_transcripts,
    process_video
)


class TestFormatTime(unittest.TestCase):
    """Test cases for format_time function"""

    def test_format_time_zero(self):
        """Test formatting 0 seconds"""
        result = format_time(0)
        self.assertEqual(result, "00:00:00,000")

    def test_format_time_seconds_only(self):
        """Test formatting seconds only"""
        result = format_time(15.5)
        self.assertEqual(result, "00:00:15,500")

    def test_format_time_minutes(self):
        """Test formatting with minutes"""
        result = format_time(65.75)  # 1 min 5.75 sec
        self.assertEqual(result, "00:01:05,750")

    def test_format_time_hours(self):
        """Test formatting with hours"""
        result = format_time(3661.123)  # 1 hour 1 min 1.123 sec
        self.assertEqual(result, "01:01:01,123")

    def test_format_time_large_value(self):
        """Test formatting large time values"""
        result = format_time(36000)  # 10 hours
        self.assertEqual(result, "10:00:00,000")

    def test_format_time_milliseconds(self):
        """Test milliseconds precision"""
        result = format_time(1.999)
        self.assertEqual(result, "00:00:01,999")


class TestCleanSrtContent(unittest.TestCase):
    """Test cases for clean_srt_content function"""

    def test_clean_basic_srt(self):
        """Test cleaning basic SRT content"""
        srt_content = """1
00:00:00,000 --> 00:00:05,000
Hello world

2
00:00:05,000 --> 00:00:10,000
This is a test
"""
        expected = "Hello world\nThis is a test"
        result = clean_srt_content(srt_content)
        self.assertEqual(result, expected)

    def test_clean_srt_with_empty_lines(self):
        """Test cleaning SRT with extra empty lines"""
        srt_content = """1
00:00:00,000 --> 00:00:05,000
Hello


2
00:00:05,000 --> 00:00:10,000
World
"""
        result = clean_srt_content(srt_content)
        self.assertNotIn("00:00:00", result)
        self.assertIn("Hello", result)
        self.assertIn("World", result)

    def test_clean_srt_removes_timestamps(self):
        """Test that timestamps are removed"""
        srt_content = """1
00:00:01,000 --> 00:00:05,000
Text here
"""
        result = clean_srt_content(srt_content)
        self.assertNotIn("-->", result)
        self.assertNotIn("00:00", result)

    def test_clean_srt_removes_indices(self):
        """Test that indices are removed"""
        srt_content = """1
00:00:00,000 --> 00:00:05,000
Text
"""
        result = clean_srt_content(srt_content)
        self.assertFalse(result.startswith("1"))

    def test_clean_empty_srt(self):
        """Test cleaning empty SRT"""
        result = clean_srt_content("")
        self.assertEqual(result, "")


class TestSaveAsDocx(unittest.TestCase):
    """Test cases for save_as_docx function"""

    def setUp(self):
        """Create temporary directory for test files"""
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Clean up temporary directory"""
        shutil.rmtree(self.temp_dir)

    def test_save_string_content(self):
        """Test saving string content as DOCX"""
        content = "Line 1\nLine 2\nLine 3"
        output_path = os.path.join(self.temp_dir, "test.docx")

        save_as_docx(content, output_path)

        self.assertTrue(os.path.exists(output_path))
        doc = Document(output_path)
        paragraphs = [p.text for p in doc.paragraphs]
        self.assertEqual(len(paragraphs), 3)
        self.assertEqual(paragraphs[0], "Line 1")

    def test_save_list_content(self):
        """Test saving list content as DOCX"""
        content = ["Paragraph 1", "Paragraph 2", "Paragraph 3"]
        output_path = os.path.join(self.temp_dir, "test_list.docx")

        save_as_docx(content, output_path)

        self.assertTrue(os.path.exists(output_path))
        doc = Document(output_path)
        paragraphs = [p.text for p in doc.paragraphs]
        self.assertEqual(len(paragraphs), 3)

    def test_save_empty_content(self):
        """Test saving empty content"""
        output_path = os.path.join(self.temp_dir, "empty.docx")
        save_as_docx("", output_path)
        self.assertTrue(os.path.exists(output_path))

    def test_save_multiline_string(self):
        """Test saving multiline string"""
        content = "First\nSecond\nThird"
        output_path = os.path.join(self.temp_dir, "multiline.docx")
        save_as_docx(content, output_path)

        doc = Document(output_path)
        self.assertEqual(len(doc.paragraphs), 3)


class TestBatchTranslate(unittest.TestCase):
    """Test cases for batch_translate function"""

    @patch('video_to_subtitles.GoogleTranslator')
    def test_batch_translate_basic(self, mock_translator_class):
        """Test basic batch translation"""
        mock_translator = MagicMock()
        mock_translator_class.return_value = mock_translator
        mock_translator.translate.return_value = "Hola |SEP| Mundo"

        text_lines = ["Hello", "World"]
        result = batch_translate(text_lines, "en", "es", batch_size=5)

        self.assertEqual(len(result), 2)
        mock_translator.translate.assert_called_once()

    @patch('video_to_subtitles.GoogleTranslator')
    def test_batch_translate_with_empty_lines(self, mock_translator_class):
        """Test translation preserves empty lines"""
        mock_translator = MagicMock()
        mock_translator_class.return_value = mock_translator
        mock_translator.translate.return_value = "Hello"

        text_lines = ["Hello", "", "World"]
        result = batch_translate(text_lines, "en", "es", batch_size=5)

        # Empty lines should be preserved
        self.assertEqual(result[1], "")

    @patch('video_to_subtitles.GoogleTranslator')
    def test_batch_translate_multiple_batches(self, mock_translator_class):
        """Test translation with multiple batches"""
        mock_translator = MagicMock()
        mock_translator_class.return_value = mock_translator
        mock_translator.translate.side_effect = [
            "Hola |SEP| Mundo",
            "Foo |SEP| Bar"
        ]

        text_lines = ["Hello", "World", "Foo", "Bar"]
        result = batch_translate(text_lines, "en", "es", batch_size=2)

        self.assertEqual(len(result), 4)
        self.assertEqual(mock_translator.translate.call_count, 2)

    @patch('video_to_subtitles.GoogleTranslator')
    def test_batch_translate_all_empty_lines(self, mock_translator_class):
        """Test translation with all empty lines"""
        mock_translator = MagicMock()
        mock_translator_class.return_value = mock_translator

        text_lines = ["", "", ""]
        result = batch_translate(text_lines, "en", "es", batch_size=5)

        self.assertEqual(result, ["", "", ""])
        mock_translator.translate.assert_not_called()

    @patch('video_to_subtitles.GoogleTranslator')
    def test_batch_translate_exception_handling(self, mock_translator_class):
        """Test handling translation exceptions"""
        mock_translator = MagicMock()
        mock_translator_class.return_value = mock_translator
        mock_translator.translate.side_effect = Exception("Translation failed")

        text_lines = ["Hello", "World"]
        result = batch_translate(text_lines, "en", "es", batch_size=5)

        # Should return original lines if translation fails
        self.assertEqual(len(result), 2)


class TestGenerateSrt(unittest.TestCase):
    """Test cases for generate_srt function"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir)

    @patch('video_to_subtitles.whisper.load_model')
    def test_generate_srt_success(self, mock_load_model):
        """Test successful SRT generation"""
        mock_model = MagicMock()
        mock_load_model.return_value = mock_model
        mock_model.transcribe.return_value = {
            "language": "en",
            "segments": [
                {"start": 0.0, "end": 5.0, "text": "Hello"},
                {"start": 5.0, "end": 10.0, "text": "World"}
            ]
        }

        video_path = os.path.join(self.temp_dir, "test.mp4")
        output_path = os.path.join(self.temp_dir, "output.srt")
        
        # Create dummy video file
        open(video_path, 'w').close()

        success, language = generate_srt(video_path, output_path)

        self.assertTrue(success)
        self.assertEqual(language, "en")
        self.assertTrue(os.path.exists(output_path))

    @patch('video_to_subtitles.whisper.load_model')
    def test_generate_srt_file_format(self, mock_load_model):
        """Test SRT file format is correct"""
        mock_model = MagicMock()
        mock_load_model.return_value = mock_model
        mock_model.transcribe.return_value = {
            "language": "en",
            "segments": [
                {"start": 0.5, "end": 5.5, "text": "Test"}
            ]
        }

        video_path = os.path.join(self.temp_dir, "test.mp4")
        output_path = os.path.join(self.temp_dir, "output.srt")
        open(video_path, 'w').close()

        generate_srt(video_path, output_path)

        with open(output_path, 'r') as f:
            content = f.read()
            self.assertIn("1", content)  # Index
            self.assertIn("00:00:00,500 --> 00:00:05,500", content)  # Timestamp
            self.assertIn("Test", content)  # Text

    @patch('video_to_subtitles.whisper.load_model')
    def test_generate_srt_exception(self, mock_load_model):
        """Test exception handling in generate_srt"""
        mock_load_model.side_effect = Exception("Model loading failed")

        video_path = os.path.join(self.temp_dir, "test.mp4")
        output_path = os.path.join(self.temp_dir, "output.srt")
        open(video_path, 'w').close()

        success, language = generate_srt(video_path, output_path)

        self.assertFalse(success)
        self.assertIsNone(language)


class TestProcessTranscripts(unittest.TestCase):
    """Test cases for process_transcripts function"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir)

    def _create_srt_file(self, path, content):
        """Helper to create SRT file"""
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)

    @patch('video_to_subtitles.batch_translate')
    def test_process_transcripts_english(self, mock_batch_translate):
        """Test processing when source language is English"""
        srt_file = os.path.join(self.temp_dir, "test.srt")
        original_txt = os.path.join(self.temp_dir, "original.txt")
        original_docx = os.path.join(self.temp_dir, "original.docx")
        english_txt = os.path.join(self.temp_dir, "english.txt")
        english_docx = os.path.join(self.temp_dir, "english.docx")
        telugu_txt = os.path.join(self.temp_dir, "telugu.txt")
        telugu_docx = os.path.join(self.temp_dir, "telugu.docx")

        srt_content = """1
00:00:00,000 --> 00:00:05,000
Hello world

2
00:00:05,000 --> 00:00:10,000
This is a test
"""
        self._create_srt_file(srt_file, srt_content)
        mock_batch_translate.return_value = ["English text"]

        success = process_transcripts(
            srt_file, original_txt, original_docx,
            english_txt, english_docx, telugu_txt, telugu_docx,
            "en"
        )

        self.assertTrue(success)
        self.assertTrue(os.path.exists(original_txt))
        self.assertTrue(os.path.exists(english_txt))
        self.assertTrue(os.path.exists(telugu_txt))

    @patch('video_to_subtitles.batch_translate')
    def test_process_transcripts_non_english(self, mock_batch_translate):
        """Test processing with non-English source"""
        srt_file = os.path.join(self.temp_dir, "test.srt")
        original_txt = os.path.join(self.temp_dir, "original.txt")
        original_docx = os.path.join(self.temp_dir, "original.docx")
        english_txt = os.path.join(self.temp_dir, "english.txt")
        english_docx = os.path.join(self.temp_dir, "english.docx")
        telugu_txt = os.path.join(self.temp_dir, "telugu.txt")
        telugu_docx = os.path.join(self.temp_dir, "telugu.docx")

        srt_content = """1
00:00:00,000 --> 00:00:05,000
Hola mundo
"""
        self._create_srt_file(srt_file, srt_content)
        mock_batch_translate.return_value = ["Hello world"]

        success = process_transcripts(
            srt_file, original_txt, original_docx,
            english_txt, english_docx, telugu_txt, telugu_docx,
            "es"
        )

        self.assertTrue(success)
        mock_batch_translate.assert_called()

    def test_process_transcripts_missing_srt(self):
        """Test handling missing SRT file"""
        srt_file = os.path.join(self.temp_dir, "nonexistent.srt")
        original_txt = os.path.join(self.temp_dir, "original.txt")
        original_docx = os.path.join(self.temp_dir, "original.docx")
        english_txt = os.path.join(self.temp_dir, "english.txt")
        english_docx = os.path.join(self.temp_dir, "english.docx")
        telugu_txt = os.path.join(self.temp_dir, "telugu.txt")
        telugu_docx = os.path.join(self.temp_dir, "telugu.docx")

        success = process_transcripts(
            srt_file, original_txt, original_docx,
            english_txt, english_docx, telugu_txt, telugu_docx,
            "en"
        )

        self.assertFalse(success)


class TestProcessVideo(unittest.TestCase):
    """Test cases for process_video function"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir)

    def test_process_video_file_not_found(self):
        """Test error when video file not found"""
        result = process_video("/nonexistent/video.mp4")
        self.assertFalse(result)

    def test_process_video_invalid_extension(self):
        """Test error with invalid file extension"""
        test_file = os.path.join(self.temp_dir, "video.avi")
        open(test_file, 'w').close()
        result = process_video(test_file)
        self.assertFalse(result)

    @patch('video_to_subtitles.generate_srt')
    @patch('video_to_subtitles.process_transcripts')
    def test_process_video_success(self, mock_process_trans, mock_generate_srt):
        """Test successful video processing"""
        mock_generate_srt.return_value = (True, "en")
        mock_process_trans.return_value = True

        test_file = os.path.join(self.temp_dir, "video.mp4")
        open(test_file, 'w').close()

        result = process_video(test_file)
        self.assertTrue(result)

    @patch('video_to_subtitles.generate_srt')
    def test_process_video_generation_fails(self, mock_generate_srt):
        """Test when SRT generation fails"""
        mock_generate_srt.return_value = (False, None)

        test_file = os.path.join(self.temp_dir, "video.mp4")
        open(test_file, 'w').close()

        result = process_video(test_file)
        self.assertFalse(result)

    @patch('video_to_subtitles.process_transcripts')
    def test_process_video_skip_transcription(self, mock_process_trans):
        """Test skip transcription mode"""
        mock_process_trans.return_value = True

        # Create test files
        video_file = os.path.join(self.temp_dir, "video.mp4")
        srt_file = os.path.join(self.temp_dir, "video_subs.srt")
        original_txt = os.path.join(self.temp_dir, "video_original.txt")

        open(video_file, 'w').close()
        
        srt_content = """1
00:00:00,000 --> 00:00:05,000
Test
"""
        with open(srt_file, 'w') as f:
            f.write(srt_content)

        with open(original_txt, 'w') as f:
            f.write("Test content")

        result = process_video(video_file, skip_transcription=True)
        self.assertTrue(result)


if __name__ == '__main__':
    unittest.main()
